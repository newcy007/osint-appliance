#!/usr/bin/env python3
"""
Generate a DOCX OSINT report from the Markdown report.

Usage:
    python reporting/generate_docx.py --target acme
"""

from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parent.parent


def generate_docx(target_name: str):
    data_dir = ROOT / "data" / target_name
    md_path = data_dir / f"{target_name}_osint_report.md"
    if not md_path.exists():
        raise SystemExit("Markdown report missing – run generate_report.py first.")

    md_text = md_path.read_text(encoding="utf-8")

    template_path = ROOT / "reporting" / "templates" / "osint_template.docx"
    if template_path.exists():
        doc = Document(template_path)
    else:
        doc = Document()

    # Simple markdown-to-heading mapping
    for line in md_text.splitlines():
        stripped = line.strip()

        if not stripped:
            doc.add_paragraph("")
            continue

        if stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=3)
        elif stripped.startswith("|") and stripped.endswith("|"):
            # Keep markdown table visually aligned – optional improvement:
            # you could parse this into real Word tables later.
            doc.add_paragraph(stripped)
        else:
            doc.add_paragraph(stripped)

    out_path = data_dir / f"{target_name}_osint_report.docx"
    doc.save(out_path)
    print(f"[*] DOCX report written to: {out_path}")


if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser()
    parser.add_argument("--target", required=True)
    args = parser.parse_args()
    generate_docx(args.target)
